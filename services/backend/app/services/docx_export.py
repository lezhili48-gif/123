from datetime import datetime
from pathlib import Path
from docx import Document


def generate_docx(project_name: str, papers: list[dict], evidence: list[dict], output_dir: Path, title: str) -> Path:
    doc = Document()
    doc.add_heading(title, 0)
    for section in ["Abstract", "Methods(检索策略)", "Results by themes", "Discussion", "Gaps", "Conclusions", "References placeholder"]:
        doc.add_heading(section, level=1)
        doc.add_paragraph("Auto-generated draft content.")

    doc.add_heading("Table 1: 纳入研究概览", level=2)
    t1 = doc.add_table(rows=1, cols=5)
    hdr = t1.rows[0].cells
    hdr[0].text = "year"
    hdr[1].text = "soil/system"
    hdr[2].text = "biochar temp/dose"
    hdr[3].text = "key outcomes"
    hdr[4].text = "cite"
    for p in papers[:30]:
        row = t1.add_row().cells
        row[0].text = str(p.get("year", ""))
        row[1].text = p.get("journal", "")
        row[2].text = "N/A"
        row[3].text = (p.get("abstract", "")[:90])
        row[4].text = f"{{CITE:{p.get('doi','NO_DOI')}}}"

    doc.add_heading("Table 2: Mechanism evidence matrix", level=2)
    t2 = doc.add_table(rows=1, cols=3)
    h2 = t2.rows[0].cells
    h2[0].text = "mechanism tag"
    h2[1].text = "outcomes"
    h2[2].text = "source"
    for e in evidence[:50]:
        r = t2.add_row().cells
        r[0].text = e.get("mechanism") or "unspecified"
        r[1].text = e.get("outcome") or e.get("claim", "")[:80]
        r[2].text = e.get("source_pointer", "")

    out = output_dir / f"{project_name}_{datetime.utcnow().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(out)
    return out
